"""
AGENTE 2 — Sintetizador y generador de reporte (summarizer.py)
================================================================
Responsabilidad única: recibir las noticias crudas del Agente 1,
seleccionar las 3 más relevantes por país, generar un resumen breve
de cada una usando Groq (Llama 3.3 70B), y producir un documento
Word (.docx) con el reporte final.
"""

import io
import time
from datetime import datetime
from typing import List, Dict

import groq
from groq import Groq
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

# ---------------------------------------------------------------------------
# 1. CONFIGURACIÓN DE GROQ
# ---------------------------------------------------------------------------

MODELO_GROQ = "llama-3.3-70b-versatile"  # buena calidad en español, rápido, free tier generoso

MAX_REINTENTOS = 3
ESPERA_BASE_SEGUNDOS = 2  # backoff exponencial: 2s, 4s, 8s
ESPERA_ENTRE_LLAMADAS_SEGUNDOS = 1.3  # ritmo entre llamadas para no chocar contra 30 RPM de Groq


def resumir_noticia(titulo: str, snippet: str, pais: str, groq_api_key: str) -> Dict:
    """
    Genera un resumen breve (2-3 frases) de una noticia usando Groq
    (Llama 3.3 70B).

    Reintenta con backoff exponencial ante errores transitorios (rate
    limit, timeouts, errores de servidor). Si todos los intentos
    fallan, retorna el snippet original como respaldo seguro.

    Returns
    -------
    {"resumen": str, "error_detalle": str | None}
    El error_detalle queda registrado (no se le muestra al usuario final
    en la UI por defecto, pero permite diagnosticar fallas reales en vez
    de ocultarlas silenciosamente).
    """
    cliente = Groq(api_key=groq_api_key)

    prompt = (
        "Eres un analista de políticas públicas. Redacta un resumen breve "
        "(máximo 3 frases, en español neutro, tono informativo y objetivo) "
        "de la siguiente noticia sobre programas de protección social en "
        f"{pais}. No inventes datos que no estén en el texto fuente.\n\n"
        f"Título: {titulo}\n"
        f"Extracto original: {snippet}\n\n"
        "Resumen:"
    )

    ultimo_error = None

    for intento in range(1, MAX_REINTENTOS + 1):
        try:
            respuesta = cliente.chat.completions.create(
                model=MODELO_GROQ,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.3,
                max_completion_tokens=200,
            )
            texto = (respuesta.choices[0].message.content or "").strip()
            if texto:
                return {"resumen": texto, "error_detalle": None}
            ultimo_error = "Groq devolvió una respuesta vacía."
            break

        except groq.RateLimitError as e:
            ultimo_error = f"Intento {intento}: RateLimitError (429) - {e}"
            if intento < MAX_REINTENTOS:
                time.sleep(ESPERA_BASE_SEGUNDOS * (2 ** (intento - 1)))
                continue

        except groq.APIStatusError as e:
            ultimo_error = f"Intento {intento}: APIStatusError {e.status_code} - {e}"
            # 4xx distintos de 429 (key inválida, modelo no encontrado, etc.)
            # no se arreglan reintentando — salimos del loop.
            if e.status_code >= 500 and intento < MAX_REINTENTOS:
                time.sleep(ESPERA_BASE_SEGUNDOS * (2 ** (intento - 1)))
                continue
            break

        except groq.APIConnectionError as e:
            ultimo_error = f"Intento {intento}: APIConnectionError - {e}"
            if intento < MAX_REINTENTOS:
                time.sleep(ESPERA_BASE_SEGUNDOS * (2 ** (intento - 1)))
                continue

        except Exception as e:
            ultimo_error = f"Intento {intento}: {type(e).__name__} - {e}"
            if intento < MAX_REINTENTOS:
                time.sleep(ESPERA_BASE_SEGUNDOS * (2 ** (intento - 1)))
                continue

    # Si llegamos aquí, todos los intentos fallaron.
    return {
        "resumen": snippet or "Resumen no disponible.",
        "error_detalle": ultimo_error,
    }


def seleccionar_top_n(noticias: List[Dict], n: int = 3) -> List[Dict]:
    """
    Selecciona las n noticias más relevantes de una lista.
    Criterio simple: las primeras n en el orden devuelto por SerpAPI,
    que ya viene ordenado por relevancia/recencia de Google News.
    Se descartan entradas con error o sin título.
    """
    validas = [
        noticia for noticia in noticias
        if "error" not in noticia and noticia.get("titulo")
    ]
    return validas[:n]


def procesar_pais(
    pais: str,
    noticias_crudas: List[Dict],
    groq_api_key: str,
    n_noticias: int = 5,
) -> Dict:
    """
    Para un país: selecciona top-N noticias y genera resumen de cada una.

    Returns
    -------
    {"pais": str, "noticias": [{"titulo", "fuente", "fecha", "link", "resumen"}],
     "sin_resultados": bool, "errores_llm": [str, ...]}
    """
    top = seleccionar_top_n(noticias_crudas, n_noticias)

    if not top:
        return {"pais": pais, "noticias": [], "sin_resultados": True, "errores_llm": []}

    procesadas = []
    errores_llm = []
    for i, noticia in enumerate(top):
        if i > 0:
            # Pequeña pausa entre llamadas consecutivas para repartir el
            # volumen dentro del límite de 30 solicitudes/minuto de Groq
            # (con 10 países x 5 noticias = 50 llamadas por ejecución).
            time.sleep(ESPERA_ENTRE_LLAMADAS_SEGUNDOS)

        resultado = resumir_noticia(
            titulo=noticia["titulo"],
            snippet=noticia.get("snippet", ""),
            pais=pais,
            groq_api_key=groq_api_key,
        )
        if resultado["error_detalle"]:
            errores_llm.append(f"{noticia['titulo'][:50]}... -> {resultado['error_detalle']}")

        procesadas.append({
            "titulo": noticia["titulo"],
            "fuente": noticia.get("fuente", "Fuente desconocida"),
            "fecha": noticia.get("fecha", ""),
            "link": noticia.get("link", ""),
            "resumen": resultado["resumen"],
        })

    return {"pais": pais, "noticias": procesadas, "sin_resultados": False, "errores_llm": errores_llm}


def procesar_todos_los_paises(
    noticias_por_pais: Dict[str, List[Dict]],
    groq_api_key: str,
    n_noticias: int = 3,
    progress_callback=None,
) -> List[Dict]:
    """
    Orquesta el procesamiento (selección + resumen) para todos los países.
    progress_callback: callback(pais_actual, indice, total) opcional.
    """
    resultado = []
    paises = list(noticias_por_pais.keys())

    for i, pais in enumerate(paises):
        if progress_callback:
            progress_callback(pais, i + 1, len(paises))

        resultado.append(
            procesar_pais(pais, noticias_por_pais[pais], groq_api_key, n_noticias)
        )

    return resultado


# ---------------------------------------------------------------------------
# 2. GENERACIÓN DEL DOCUMENTO WORD
# ---------------------------------------------------------------------------

COLOR_TITULO = RGBColor(0x1F, 0x3A, 0x5F)   # azul oscuro institucional
COLOR_PAIS = RGBColor(0x2E, 0x75, 0xB6)     # azul medio
COLOR_GRIS = RGBColor(0x59, 0x59, 0x59)     # gris para metadatos


def _configurar_estilos(doc: Document):
    """Define fuente Arial por defecto y tamaños consistentes."""
    estilo_normal = doc.styles["Normal"]
    estilo_normal.font.name = "Arial"
    estilo_normal.font.size = Pt(11)
    # Asegurar fuente también para texto de Asia oriental (consistencia en Word)
    rpr = estilo_normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), "Arial")


def _agregar_portada(doc: Document, fecha_str: str):
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run("Resumen de Noticias\nProgramas de Protección Social en Centroamérica y el Caribe")
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = COLOR_TITULO

    subtitulo = doc.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitulo.add_run(f"Reporte generado el {fecha_str}")
    run2.font.size = Pt(12)
    run2.font.italic = True
    run2.font.color.rgb = COLOR_GRIS

    nota = doc.add_paragraph()
    nota.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = nota.add_run(
        "Países incluidos: Costa Rica, Cuba, El Salvador, Guatemala, Haití, "
        "Honduras, México, Nicaragua, Panamá y República Dominicana."
    )
    run3.font.size = Pt(10)
    run3.font.color.rgb = COLOR_GRIS

    doc.add_paragraph()  # espacio


def _agregar_pais(doc: Document, datos_pais: Dict):
    pais = datos_pais["pais"]

    encabezado = doc.add_heading(level=1)
    run = encabezado.add_run(pais)
    run.font.color.rgb = COLOR_PAIS
    run.font.name = "Arial"

    if datos_pais["sin_resultados"]:
        p = doc.add_paragraph()
        run_vacio = p.add_run(
            "No se encontraron noticias relevantes sobre programas de "
            "protección social en las últimas 24 horas para este país."
        )
        run_vacio.font.italic = True
        run_vacio.font.color.rgb = COLOR_GRIS
        doc.add_paragraph()
        return

    for idx, noticia in enumerate(datos_pais["noticias"], start=1):
        sub = doc.add_heading(level=2)
        sub_run = sub.add_run(f"{idx}. {noticia['titulo']}")
        sub_run.font.size = Pt(13)
        sub_run.font.name = "Arial"
        sub_run.font.color.rgb = COLOR_TITULO

        meta = doc.add_paragraph()
        meta_run = meta.add_run(f"Fuente: {noticia['fuente']}  |  Fecha: {noticia['fecha']}")
        meta_run.font.size = Pt(9)
        meta_run.font.italic = True
        meta_run.font.color.rgb = COLOR_GRIS

        resumen = doc.add_paragraph()
        resumen.add_run(noticia["resumen"]).font.size = Pt(11)

        if noticia.get("link"):
            link_p = doc.add_paragraph()
            link_run = link_p.add_run(f"Enlace: {noticia['link']}")
            link_run.font.size = Pt(9)
            link_run.font.color.rgb = RGBColor(0x10, 0x6E, 0xBE)

        doc.add_paragraph()  # espacio entre noticias

    # línea divisoria simple entre países (borde inferior de un párrafo)
    divisor = doc.add_paragraph()
    pPr = divisor._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pPr.makeelement(qn("w:bottom"), {
        qn("w:val"): "single", qn("w:sz"): "6", qn("w:space"): "1", qn("w:color"): "2E75B6"
    })
    pBdr.append(bottom)
    pPr.append(pBdr)


def generar_documento_word(reportes_por_pais: List[Dict]) -> io.BytesIO:
    """
    Genera el .docx final en memoria (BytesIO) — ideal para servirlo
    directamente como descarga desde Streamlit sin tocar el disco.
    """
    doc = Document()
    _configurar_estilos(doc)

    MESES_ES = {
        1: "enero", 2: "febrero", 3: "marzo", 4: "abril", 5: "mayo", 6: "junio",
        7: "julio", 8: "agosto", 9: "septiembre", 10: "octubre", 11: "noviembre", 12: "diciembre",
    }
    ahora = datetime.now()
    fecha_str = f"{ahora.day} de {MESES_ES[ahora.month]} de {ahora.year}, {ahora.strftime('%H:%M')}"

    _agregar_portada(doc, fecha_str)

    for datos_pais in reportes_por_pais:
        _agregar_pais(doc, datos_pais)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
